import os
import io
import json

from flask import Flask, request, jsonify
from dotenv import load_dotenv

import pdfplumber
from docx import Document
from PIL import Image
import pytesseract

from google import genai
from google.genai import types

from flask_cors import CORS

# ------------------------------------------------
# Load environment variables
# ------------------------------------------------
load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

if not GEMINI_API_KEY:
    raise ValueError("GEMINI_API_KEY not found in .env")

# If Tesseract is not in PATH:
# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# ------------------------------------------------
# Initialize Gemini client
# ------------------------------------------------
client = genai.Client(api_key=GEMINI_API_KEY)

# ------------------------------------------------
# Flask app
# ------------------------------------------------
app = Flask(__name__)
CORS(app)

app.config["MAX_CONTENT_LENGTH"] = 10 * 1024 * 1024  # 10 MB max upload

# ------------------------------------------------
# Helpers: extract text from files
# ------------------------------------------------
def extract_text_from_pdf(file_bytes: bytes) -> str:
    text = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text.append(page_text)
    return "\n".join(text).strip()

def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs]
    return "\n".join(paragraphs).strip()

def extract_text_from_image(file_bytes: bytes) -> str:
    image = Image.open(io.BytesIO(file_bytes))
    text = pytesseract.image_to_string(image)
    return text.strip()

def extract_text(file_storage) -> str:
    filename = file_storage.filename.lower()
    file_bytes = file_storage.read()

    if filename.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif filename.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    elif filename.endswith((".png", ".jpg", ".jpeg")):
        return extract_text_from_image(file_bytes)
    else:
        return extract_text_from_pdf(file_bytes)

# ------------------------------------------------
# Convert resume text → JSON using Gemini
# ------------------------------------------------
def parse_resume_with_gemini(raw_text: str) -> dict:

    system_prompt = """
You are an expert resume parser.
Extract structured data and return ONLY valid JSON.
"""

    schema_prompt = """
Use this JSON format:

{
  "name": "string or null",
  "email": "string or null",
  "phone": "string or null",
  "location": "string or null",
  "summary": "string or null",
  "skills": [],
  "experience": [],
  "education": [],
  "projects": [],
  "social_links": {
    "linkedin": null,
    "github": null,
    "portfolio": null
  },
  "certifications": [],
  "extras": {
    "raw_text_excerpt": "string or null"
  }
}
"""

    user_prompt = f"""
Here is the resume text:

\"\"\"{raw_text}\"\"\"

Return ONLY JSON. 
Fill extras.raw_text_excerpt with a short excerpt.
"""

    # ⭐ FIXED Gemini API call
    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=[
            system_prompt,
            schema_prompt,
            user_prompt,
        ],
        config=types.GenerateContentConfig(response_modalities=["TEXT"])
    )

    raw_output = response.text.strip()

    # Remove accidental markdown
    if raw_output.startswith("```"):
        raw_output = raw_output.strip("`")
        raw_output = raw_output.replace("json", "", 1).strip()

    try:
        return json.loads(raw_output)
    except:
        return {
            "error": "JSON parsing failed",
            "raw_output": raw_output
        }

# ------------------------------------------------
# API Route
# ------------------------------------------------
@app.route("/extract-resume", methods=["POST"])
def extract_resume():
    if "file" not in request.files:
        return jsonify({"error": "No file in request"}), 400

    file = request.files["file"]

    if file.filename == "":
        return jsonify({"error": "No selected file"}), 400

    try:
        raw_text = extract_text(file)

        if not raw_text.strip():
            return jsonify({"error": "Could not extract text"}), 400

        parsed = parse_resume_with_gemini(raw_text)

        return jsonify(parsed), 200

    except Exception as e:
        print("Error:", e)
        return jsonify({"error": "Server error", "details": str(e)}), 500

@app.route("/", methods=["GET"])
def main():
    return jsonify({"message": "Resume Parser Server Running"}), 200

# ------------------------------------------------
# Run the server
# ------------------------------------------------
if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)



# import os
# import io
# import json

# from flask import Flask, request, jsonify
# from dotenv import load_dotenv

# import pdfplumber
# from docx import Document
# from PIL import Image
# import pytesseract

# from google import genai
# from google.genai import types

# from flask_cors import CORS


# # ------------------------------------------------
# # Load environment variables
# # ------------------------------------------------
# load_dotenv()
# GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

# if not GEMINI_API_KEY:
#     raise ValueError("GEMINI_API_KEY not found in .env")

# # If Tesseract is not in PATH, set it manually (Windows example)
# # pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# # ------------------------------------------------
# # Initialize Gemini client
# # ------------------------------------------------
# client = genai.Client(api_key=GEMINI_API_KEY)

# # ------------------------------------------------
# # Flask app
# # ------------------------------------------------
# app = Flask(__name__)

# CORS(app)

# # Limit upload size if you want (e.g. 10 MB)
# app.config["MAX_CONTENT_LENGTH"] = 10 * 1024 * 1024  # 10 MB


# # ------------------------------------------------
# # Helpers: extract text from different file types
# # ------------------------------------------------
# def extract_text_from_pdf(file_bytes: bytes) -> str:
#     text = []
#     with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
#         for page in pdf.pages:
#             page_text = page.extract_text() or ""
#             text.append(page_text)
#     return "\n".join(text).strip()


# def extract_text_from_docx(file_bytes: bytes) -> str:
#     doc = Document(io.BytesIO(file_bytes))
#     paragraphs = [p.text for p in doc.paragraphs]
#     return "\n".join(paragraphs).strip()


# def extract_text_from_image(file_bytes: bytes) -> str:
#     image = Image.open(io.BytesIO(file_bytes))
#     text = pytesseract.image_to_string(image)
#     return text.strip()


# def extract_text(file_storage) -> str:
#     """
#     Detects file type from filename / mimetype
#     and routes to correct extractor.
#     """
#     filename = file_storage.filename.lower()
#     file_bytes = file_storage.read()

#     if filename.endswith(".pdf"):
#         return extract_text_from_pdf(file_bytes)
#     elif filename.endswith(".docx"):
#         return extract_text_from_docx(file_bytes)
#     elif filename.endswith((".png", ".jpg", ".jpeg")):
#         return extract_text_from_image(file_bytes)
#     else:
#         # fallback: try PDF first, then image OCR
#         return extract_text_from_pdf(file_bytes)


# # ------------------------------------------------
# # Helper: call Gemini to convert resume text → JSON
# # ------------------------------------------------
# def parse_resume_with_gemini(raw_text: str) -> dict:
#     """
#     Sends raw resume text to Gemini and asks for STRICT JSON.
#     """

#     system_prompt = """
# You are an expert resume parser. 
# You will receive the full text of a candidate resume.

# Your job:
# 1. Extract as much structured information as possible.
# 2. Return ONLY valid JSON. No explanation, no markdown, no comments.
# 3. If any field is unknown, use null or empty list.

# Use exactly this JSON structure:

# {
#   "name": "string or null",
#   "email": "string or null",
#   "phone": "string or null",
#   "location": "string or null",
#   "summary": "string or null",
#   "skills": ["skill1", "skill2"],
#   "experience": [
#     {
#       "job_title": "string or null",
#       "company": "string or null",
#       "location": "string or null",
#       "start_date": "string or null",
#       "end_date": "string or null",
#       "description": "string or null"
#     }
#   ],
#   "education": [
#     {
#       "degree": "string or null",
#       "field": "string or null",
#       "institution": "string or null",
#       "start_date": "string or null",
#       "end_date": "string or null"
#     }
#   ],
#   "projects": [
#     {
#       "name": "string or null",
#       "description": "string or null",
#       "tech_stack": ["string"],
#       "link": "string or null"
#     }
#   ],
#   "social_links": {
#     "linkedin": "string or null",
#     "github": "string or null",
#     "portfolio": "string or null"
#   },
#   "certifications": [
#     {
#       "name": "string or null",
#       "issuer": "string or null",
#       "year": "string or null"
#     }
#   ],
#   "extras": {
#     "raw_text_excerpt": "string or null"
#   }
# }

# Remember:
# - Return ONLY JSON.
# - Do not wrap in ```json or any markdown.
# - Do not include any extra keys.
# """

#     excerpt = raw_text[:1000]  # keep for extras

#     user_prompt = f"""
# Here is the resume text:

# \"\"\"{raw_text}\"\"\"

# Also, fill extras.raw_text_excerpt with a short excerpt from the text (max 300 chars).
# """

#     response = client.models.generate_content(
#         model="gemini-2.5-flash",
#         contents=[
#             types.Part.from_text(system_prompt),
#             types.Part.from_text(user_prompt),
#         ],
#         config=types.GenerateContentConfig(
#             response_modalities=["TEXT"]
#         )
#     )

#     # Sometimes model returns plain text, sometimes in candidates/parts
#     raw_output = response.text or ""
#     raw_output = raw_output.strip()

#     # Remove code fences if model still returns them
#     if raw_output.startswith("```"):
#         raw_output = raw_output.strip("`")
#         # remove possible "json" after opening ```
#         raw_output = raw_output.replace("json", "", 1).strip()

#     try:
#         data = json.loads(raw_output)
#     except json.JSONDecodeError:
#         # fallback: return minimal structure
#         data = {
#             "name": None,
#             "email": None,
#             "phone": None,
#             "location": None,
#             "summary": None,
#             "skills": [],
#             "experience": [],
#             "education": [],
#             "projects": [],
#             "social_links": {
#                 "linkedin": None,
#                 "github": None,
#                 "portfolio": None
#             },
#             "certifications": [],
#             "extras": {
#                 "raw_text_excerpt": excerpt
#             }
#         }

#     return data


# # ------------------------------------------------
# # API Route
# # ------------------------------------------------
# @app.route("/extract-resume", methods=["POST"])
# def extract_resume():
#     if "file" not in request.files:
#         return jsonify({"error": "No file part in request"}), 400

#     file = request.files["file"]

#     if file.filename == "":
#         return jsonify({"error": "No selected file"}), 400

#     try:
#         # 1. Extract raw text from PDF / DOCX / image
#         raw_text = extract_text(file)

#         if not raw_text.strip():
#             return jsonify({"error": "Unable to extract text from file"}), 400

#         # 2. Ask Gemini to convert to structured JSON
#         parsed = parse_resume_with_gemini(raw_text)

#         return jsonify(parsed), 200

#     except Exception as e:
#         print("Error:", e)
#         return jsonify({"error": "Internal server error", "details": str(e)}), 500

# @app.route("/", methods=["GET"])
# def main():
#     return jsonify({"success": "Server is running at a speed of 400kmph"})


# # ------------------------------------------------
# # Run the server
# # ------------------------------------------------
# if __name__ == "__main__":
#     # For local dev
#     app.run(host="0.0.0.0", port=5000, debug=True)
